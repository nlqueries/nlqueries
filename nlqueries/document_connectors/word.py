"""
nlqueries.document_connectors.word
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Microsoft Word (.docx) document connector using ``python-docx`` for paragraph
extraction and a built-in recursive character chunker.

Chunking strategy: heading-based sections.  The connector walks ``doc.paragraphs``
and treats every ``Heading 1`` / ``Heading 2`` paragraph as a section boundary.
Text accumulated under a heading is flushed as one or more chunks when the next
heading is encountered or the document ends.  If the accumulated section text
exceeds 1 200 characters it is split with ``RecursiveCharacterTextSplitter``;
otherwise the whole section becomes a single chunk.

Requires the ``docs`` optional dependency group:
    pip install "nlqueries-core[docs]"
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from nlqueries.document_connectors.base import DocumentChunk, DocumentConnector
from nlqueries.document_connectors.chunker import RecursiveCharacterTextSplitter

_HEADING_STYLES = {"Heading 1", "Heading 2"}
_SPLIT_THRESHOLD = 1_200
_CHUNK_SIZE = 800
_CHUNK_OVERLAP = 100


def _make_chunk_id(source_id: str, chunk_index: int) -> str:
    """Deterministic 16-char hex ID: sha256(source_id:None:chunk_index)[:16]."""
    raw = f"{source_id}:None:{chunk_index}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


class WordConnector(DocumentConnector):
    """Extract and chunk text from Word (.docx) files using python-docx.

    Word files have no concept of pages without rendering, so ``page_number``
    is always ``None``.  Sections are delimited by Heading 1 / Heading 2
    paragraphs; documents without headings are treated as a single section.
    """

    def ingest(self, source: str | Path, source_id: str) -> list[DocumentChunk]:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for WordConnector. "
                "Install it with: pip install 'nlqueries-core[docs]'"
            ) from exc

        source_path = Path(source)
        doc = docx.Document(str(source_path))

        # Collect (heading_text, body_text) pairs by walking paragraphs.
        sections: list[tuple[str, str]] = []
        current_heading = "untitled"
        current_lines: list[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if style_name in _HEADING_STYLES:
                # Flush the previous section before starting a new one.
                body = "\n".join(current_lines).strip()
                if body:
                    sections.append((current_heading, body))
                current_heading = text or "untitled"
                current_lines = []
            else:
                if text:
                    current_lines.append(text)

        # Flush the last section.
        body = "\n".join(current_lines).strip()
        if body:
            sections.append((current_heading, body))

        # If the document had no headings at all, treat everything as one section.
        if not sections:
            all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            if all_text:
                sections = [("untitled", all_text)]

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=_CHUNK_SIZE,
            chunk_overlap=_CHUNK_OVERLAP,
        )

        chunks: list[DocumentChunk] = []
        global_chunk_index = 0

        for heading, section_text in sections:
            if len(section_text) > _SPLIT_THRESHOLD:
                sub_texts = splitter.split_text(section_text)
            else:
                sub_texts = [section_text]

            for sub_text in sub_texts:
                chunk_id = _make_chunk_id(source_id, global_chunk_index)
                chunks.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        source_id=source_id,
                        source_name=source_path.name,
                        page_number=None,
                        chunk_index=global_chunk_index,
                        text=sub_text,
                        metadata={
                            "connector": "word",
                            "file_path": str(source_path),
                            "section_heading": heading,
                        },
                    )
                )
                global_chunk_index += 1

        return chunks

    def supports(self, source: str | Path) -> bool:
        return Path(source).suffix.lower() == ".docx"
